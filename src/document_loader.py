# src/document_loader.py
"""
Document Loader — handles PDF, DOCX, and TXT files.
Extracts raw text + metadata (source, page, filetype) from each file.
Metadata is attached at load time so citations work later in the RAG chain.
"""

import fitz                          # PyMuPDF — for PDF parsing
from docx import Document as DocxDocument   # python-docx — for DOCX parsing
from pathlib import Path
from typing import List
from langchain.schema import Document       # LangChain's Document wrapper
from src.config import SUPPORTED_EXTENSIONS


# ─────────────────────────────────────────────────────────────
# MAIN LOADER — dispatches to the right parser by file type
# ─────────────────────────────────────────────────────────────

def load_document(file_path: str) -> List[Document]:
    """
    Load a single document and return a list of LangChain Document objects.
    Each Document contains:
      - page_content : the raw text of that page/section
      - metadata     : {source, page, filetype, filename}

    Args:
        file_path: Absolute or relative path to the file.

    Returns:
        List of Document objects, one per page (PDF) or section (DOCX/TXT).

    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    # Guard: file must exist
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Guard: file type must be supported
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {SUPPORTED_EXTENSIONS}"
        )

    # Dispatch to the right parser
    if ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".docx":
        return _load_docx(path)
    elif ext == ".txt":
        return _load_txt(path)


def load_documents(file_paths: List[str]) -> List[Document]:
    """
    Load multiple documents at once.
    Skips files that fail to load and logs the error.

    Args:
        file_paths: List of file paths to load.

    Returns:
        Combined list of Document objects from all files.
    """
    all_documents = []

    for file_path in file_paths:
        try:
            docs = load_document(file_path)
            all_documents.extend(docs)
            print(f"  ✅ Loaded: {Path(file_path).name} "
                  f"({len(docs)} page(s)/section(s))")
        except Exception as e:
            print(f"  ❌ Failed to load {file_path}: {e}")

    print(f"\n  📄 Total sections loaded: {len(all_documents)}")
    return all_documents


# ─────────────────────────────────────────────────────────────
# PDF PARSER — one Document per page
# ─────────────────────────────────────────────────────────────

def _load_pdf(path: Path) -> List[Document]:
    """
    Parse a PDF file using PyMuPDF.
    Returns one Document per page with page number in metadata.
    """
    documents = []

    # Open PDF with PyMuPDF
    pdf = fitz.open(str(path))

    for page_num in range(len(pdf)):
        page = pdf[page_num]

        # Extract plain text from the page
        text = page.get_text("text").strip()

        # Skip empty pages (scanned images, blank pages, etc.)
        if not text:
            continue

        # Build the Document with rich metadata for citations
        doc = Document(
            page_content=text,
            metadata={
                "source"   : str(path),        # full file path
                "filename" : path.name,         # e.g. "research_paper.pdf"
                "page"     : page_num + 1,      # human-readable (1-indexed)
                "filetype" : "pdf",
                "total_pages": len(pdf),
            }
        )
        documents.append(doc)

    pdf.close()
    return documents


# ─────────────────────────────────────────────────────────────
# DOCX PARSER — one Document per paragraph group
# ─────────────────────────────────────────────────────────────

def _load_docx(path: Path) -> List[Document]:
    """
    Parse a DOCX file using python-docx.
    Groups paragraphs into sections of ~500 chars to mimic page structure.
    Returns one Document per section.
    """
    docx_file = DocxDocument(str(path))
    documents = []

    # Collect all non-empty paragraphs
    paragraphs = [
        p.text.strip()
        for p in docx_file.paragraphs
        if p.text.strip()
    ]

    if not paragraphs:
        return documents

    # Group paragraphs into sections (~500 chars each)
    # This prevents single paragraphs from becoming too tiny to be useful
    section_text  = ""
    section_index = 1

    for para in paragraphs:
        section_text += para + "\n"

        if len(section_text) >= 500:
            doc = Document(
                page_content=section_text.strip(),
                metadata={
                    "source"   : str(path),
                    "filename" : path.name,
                    "page"     : section_index,   # section number
                    "filetype" : "docx",
                    "total_pages": None,          # DOCX has no fixed page count
                }
            )
            documents.append(doc)
            section_text  = ""
            section_index += 1

    # Don't forget the last section (may be smaller than 500 chars)
    if section_text.strip():
        doc = Document(
            page_content=section_text.strip(),
            metadata={
                "source"   : str(path),
                "filename" : path.name,
                "page"     : section_index,
                "filetype" : "docx",
                "total_pages": None,
            }
        )
        documents.append(doc)

    return documents


# ─────────────────────────────────────────────────────────────
# TXT PARSER — one Document per file
# ─────────────────────────────────────────────────────────────

def _load_txt(path: Path) -> List[Document]:
    """
    Parse a plain text file.
    Returns a single Document for the entire file.
    The text splitter in the next stage will chunk it further.
    """
    text = path.read_text(encoding="utf-8", errors="ignore").strip()

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source"   : str(path),
                "filename" : path.name,
                "page"     : 1,
                "filetype" : "txt",
                "total_pages": 1,
            }
        )
    ]